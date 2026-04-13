"""
Deterministic DOCX rendering for synthesized executive Markdown.

Reads only the file at md_path (typically executive_synthesized.md). Optional ## headings
become section titles; prose without headings is rendered under a single body block.
No stitching from other report sources.
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path


def _load_text(path: str) -> str:
    p = Path(path)
    if not p.exists():
        return ""
    return p.read_text(encoding="utf-8", errors="ignore")


def _apply_styles(doc) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, RGBColor

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for name, size, bold in (
        ("SA Title", 30, True),
        ("SA Section Header", 17, True),
        ("SA Subheader", 13, True),
        ("SA Body", 11, False),
    ):
        if name in styles:
            st = styles[name]
        else:
            st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0, 0, 0)


def _add_paragraphs_for_body(doc, text: str) -> None:
    """Add body text: split paragraphs on blank lines; bullets as list."""
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    block = (text or "").strip()
    if not block:
        return
    for para in re.split(r"\n\s*\n+", block):
        para = para.strip()
        if not para:
            continue
        if para.startswith("- ") or para.startswith("* "):
            for ln in para.split("\n"):
                ln = ln.strip()
                if ln.startswith(("- ", "* ")):
                    doc.add_paragraph(ln[2:].strip(), style="List Bullet")
                elif ln:
                    doc.add_paragraph(ln, style="SA Body")
        else:
            p = doc.add_paragraph(style="SA Body")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.add_run(para)


def _add_cover(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = doc.add_paragraph("AI Site Auditor", style="SA Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph("Executive Report", style="SA Section Header")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    row = doc.add_paragraph(style="SA Body")
    row.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.add_run(f"Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}").font.size = Pt(11)
    row2 = doc.add_paragraph(style="SA Body")
    row2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row2.add_run("Confidential — internal or client use").font.size = Pt(11)
    doc.add_page_break()


def build_executive_docx(md_path: str, output_path: str) -> None:
    """
    Render synthesized executive Markdown to DOCX.

    Reads only the file at md_path (executive_synthesized.md from the build pipeline).
    """
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required for DOCX build") from exc

    md = _load_text(md_path)
    if not md.strip():
        raise ValueError(f"Synthesized markdown is empty or missing: {md_path}")

    doc = Document()
    _apply_styles(doc)
    _add_cover(doc)

    # Split on ## headings (level-2 only)
    parts = re.split(r"(?m)^##\s+(.+?)\s*$", md)
    preamble = parts[0].strip() if parts else ""
    if preamble:
        doc.add_paragraph("Overview", style="SA Subheader")
        _add_paragraphs_for_body(doc, preamble)

    i = 1
    while i + 1 < len(parts):
        title = parts[i].strip()
        body = parts[i + 1].strip()
        doc.add_paragraph(title, style="SA Section Header")
        _add_paragraphs_for_body(doc, body)
        i += 2

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def build_pptx_from_json(json_path: str, template_path: str) -> None:
    """Future extension placeholder."""
    _ = (json_path, template_path)
    raise NotImplementedError("PPTX builder is reserved for a future phase.")
